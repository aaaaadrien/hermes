"""
hermes_tools.py
===================
Module partagé contenant les outils (météo, wiki, change).
Importé par hermes-cli.py et hermes-web.py.
Important : pas de - mais des _ dans le fichier pour importer le module !

Utilisation :
    from hermes_tools import CATALOGUE_OUTILS, outils_actifs, executer_outil
"""

import configparser
import math
import re
import mimetypes
import os
import subprocess
import tempfile
from pathlib import Path
from datetime import datetime
from typing import Optional
import urllib.parse

import requests
from bs4 import BeautifulSoup
from ddgs import DDGS

# Récupère le bulletin météo complet pour une ville donnée.
def outil_meteo(ville: str) -> str:
    """
      1. Géolocalisation via l'API open-meteo geocoding
      2. Récupération des données météo (température, humidité, vent, état du ciel)
      3. Traduction du code météo WMO en texte lisible
    """
    try:
        geo_url = (
            f"https://geocoding-api.open-meteo.com/v1/search"
            f"?name={ville}&count=1&language=fr&format=json"
        )
        geo = requests.get(geo_url, timeout=5).json()
        if not geo.get("results"):
            return f"Désolé, la ville « {ville} » est introuvable."

        lieu = geo["results"][0]
        lat, lon = lieu["latitude"], lieu["longitude"]

        meteo_url = (
            f"https://api.open-meteo.com/v1/forecast"
            f"?latitude={lat}&longitude={lon}"
            f"&current=temperature_2m,relative_humidity_2m,weather_code,wind_speed_10m"
            f"&timezone=auto"
        )
        res     = requests.get(meteo_url, timeout=5).json()
        donnees = res["current"]

        traduction_temps = {
            0:  "Ciel dégagé ☀️",    1:  "Principalement dégagé 🌤️",
            2:  "Partiellement nuageux ⛅", 3: "Couvert ☁️",
            45: "Brouillard 🌫️",     48: "Brouillard givrant 🌫️❄️",
            51: "Bruine légère 🌦️",  61: "Pluie légère 🌧️",
            63: "Pluie modérée 🌧️",  65: "Pluie forte 🌧️💧",
            71: "Chute de neige légère 🌨️", 95: "Orage ⛈️",
        }
        etat = traduction_temps.get(donnees["weather_code"], "Conditions variables")

        return (
            f"Météo à {lieu['name']} ({lieu.get('country', '')}) : {etat}\n"
            f"  🌡 Température : {donnees['temperature_2m']}°C\n"
            f"  💧 Humidité    : {donnees['relative_humidity_2m']}%\n"
            f"  💨 Vent        : {donnees['wind_speed_10m']} km/h"
        )

    except requests.exceptions.Timeout:
        return "⚠️ Délai d'attente dépassé pour l'API météo."
    except Exception as e:
        return f"⚠️ Impossible de récupérer la météo ({e})."

# Retourne le résumé Wikipédia (version française) du sujet demandé.
def outil_wiki(sujet: str) -> str:
    """
    Utilise l'API REST de Wikipédia pour obtenir l'extrait de la page.
    """

    # Fix sinon bloqué
    headers = {
        'User-Agent': 'Linuxtricks/1.0'
    }
    
    try:
        url  = f"https://fr.wikipedia.org/api/rest_v1/page/summary/{sujet.replace(' ', '_')}"
        res  = requests.get(url, headers=headers, timeout=5).json()
        texte = res.get("extract")
        if not texte:
            return f"Aucune page Wikipédia trouvée pour « {sujet} »."
        return texte

    except requests.exceptions.Timeout:
        return "⚠️ Délai d'attente dépassé pour Wikipédia."
    except Exception as e:
        return f"⚠️ Erreur lors de la recherche Wikipédia ({e})."
        
# Retourne la page Wikipédia (version française) du sujet demandé.
def outil_wiki_full(sujet: str) -> str:
    """
    Utilise l'Action API de Wikipédia pour obtenir l'entièreté du contenu d'une page.
    """
    url = "https://fr.wikipedia.org/w/api.php"
    
    # Fix sinon bloqué
    headers = {
        'User-Agent': 'Linuxtricks/1.0'
    }
    
    params = {
        "action": "query",
        "prop": "extracts",
        "explaintext": True,  # Récupère le texte brut sans HTML
        "titles": sujet,
        "format": "json",
        "redirects": 1
    }

    try:
        response = requests.get(url, params=params, headers=headers, timeout=5)
        # Vérifie si la requête a réussi avant de tenter le décodage JSON
        response.raise_for_status()
        
        data = response.json()
        pages = data.get("query", {}).get("pages", {})
        
        # Récupération de la première page trouvée
        page_id = next(iter(pages))
        page_data = pages[page_id]

        if "missing" in page_data or page_id == "-1":
            return f"Aucune page Wikipédia trouvée pour « {sujet} »."

        texte = page_data.get("extract")
        if not texte:
            return f"Le contenu de la page « {sujet} » est vide."
            
        return texte

    except requests.exceptions.Timeout:
        return "⚠️ Délai d'attente dépassé pour Wikipédia."
    except requests.exceptions.HTTPError as e:
        return f"⚠️ Erreur HTTP (Accès refusé ou page inexistante) : {e}"
    except Exception as e:
        return f"⚠️ Erreur lors de la recherche Wikipédia ({e})."
        

# Convertit une somme d'une devise vers une autre en temps réel.
def outil_argent(montant: float, de_monnaie: str, vers_monnaie: str) -> str:
    """
    Utilise l'API open.er-api.com (gratuite, sans clé).

    Paramètres :
      montant      : valeur numérique à convertir
      de_monnaie   : code ISO 4217 source (ex : EUR)
      vers_monnaie : code ISO 4217 cible  (ex : USD)
    """
    try:
        url = f"https://open.er-api.com/v6/latest/{de_monnaie.upper()}"
        res = requests.get(url, timeout=5).json()

        if res.get("result") != "success":
            return f"⚠️ Devise source « {de_monnaie} » non reconnue."

        taux = res["rates"].get(vers_monnaie.upper())
        if taux is None:
            return f"⚠️ Devise cible « {vers_monnaie} » non reconnue."

        resultat = montant * taux
        return (
            f"💱 {montant} {de_monnaie.upper()} = {resultat:.2f} {vers_monnaie.upper()}\n"
            f"  Taux : 1 {de_monnaie.upper()} = {taux:.4f} {vers_monnaie.upper()}"
        )

    except requests.exceptions.Timeout:
        return "⚠️ Délai d'attente dépassé pour l'API de change."
    except Exception as e:
        return f"⚠️ Erreur de conversion monétaire ({e})."


# Effectue une recherche web via DuckDuckGo et retourne les N premiers résultats.
def outil_duckduckgo(query: str, num_results: int = 5) -> str:
    """
    Utilise DuckDuckGo (ddgs) pour rechercher sur le web.
    Retourne titres, URLs et extraits pour chaque résultat.
    """
    num_results = min(int(num_results), 10)
    try:
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=num_results))
        if not results:
            return "Aucun résultat trouvé."
        lines = [f"**Résultats pour « {query} »**\n"]
        for i, r in enumerate(results, 1):
            lines.append(
                f"{i}. **{r.get('title', 'Sans titre')}**\n"
                f"   URL : {r.get('href', '')}\n"
                f"   {r.get('body', '')}\n"
            )
        return "\n".join(lines)
    except Exception as e:
        return f"⚠️ Erreur lors de la recherche : {e}"


# Télécharge et extrait le texte principal d'une page web.
def outil_recup_page(url: str, max_chars: int = 1024000) -> str:
    """
    Récupère une page web via requests + BeautifulSoup.
    Supprime les balises inutiles (nav, scripts, pubs…) et nettoie le texte.
    """
    max_chars = min(int(max_chars), 1024000)
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
        )
    }
    try:
        resp = requests.get(url, headers=headers, timeout=10)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
        text = re.sub(r"\n{3,}", "\n\n", text)
        if len(text) > max_chars:
            text = text[:max_chars] + f"\n\n[… contenu tronqué à {max_chars} caractères]"
        return text or "Page vide ou contenu non extractible."
    except requests.exceptions.Timeout:
        return "⚠️ Délai d'attente dépassé lors du chargement de la page."
    except Exception as e:
        return f"⚠️ Erreur lors du chargement de la page : {e}"


# Extensions vidéo pour conversion en audio (via ffmpeg systeme) avant envoi à whisper.cpp
_EXTENSIONS_VIDEO = {".mp4", ".mkv", ".mov", ".avi", ".webm"}


def _extraire_audio_ffmpeg(donnees: bytes, nom_fichier: str, ignorer_debut_secondes: float = 0,
                            ignorer_fin_secondes: float = 0) -> bytes:
    """
    Extrait la piste audio d'un fichier vidéo et la convertit en WAV mono 16 kHz
    (format standard attendu par whisper.cpp) via ffmpeg.
    Note : on supprime les silences et accélère un peu l'audio (plus rapide à traiter)

    ignorer_debut_secondes : si > 0, coupe ce nombre de secondes au début du fichier
                             (pour pas traiter une intro de live)
    ignorer_fin_secondes   : si > 0, coupe ce nombre de secondes à la fin du fichier

    Nécessite que le binaire `ffmpeg` soit installé et accessible dans le PATH.
    """
    suffixe_in = Path(nom_fichier).suffix or ".mp4"
    with tempfile.NamedTemporaryFile(suffix=suffixe_in, delete=False) as f_in:
        f_in.write(donnees)
        chemin_in = f_in.name
    chemin_out = chemin_in + ".wav"

    try:
        commande = ["ffmpeg", "-y"]
        if ignorer_debut_secondes > 0:
            commande += ["-ss", str(ignorer_debut_secondes)]
        commande += ["-i", chemin_in]
        if ignorer_fin_secondes > 0:
            duree_totale = _duree_fichier_secondes(chemin_in)
            duree_utile = max(duree_totale - ignorer_debut_secondes - ignorer_fin_secondes, 1)
            commande += ["-t", str(duree_utile)]
        commande += [
            "-vn",              # pas de flux vidéo
            "-ac", "1",         # mono car plus leger
            "-ar", "16000",     # 16 kHz
            "-af", "silenceremove=stop_periods=-1:stop_duration=1:stop_threshold=-40dB,atempo=1.2", # Supprimer les blancs et accélérer pour avoir moins long à traiter
            "-f", "wav",
            chemin_out,
        ]
        resultat = subprocess.run(commande, capture_output=True, timeout=900)
        if resultat.returncode != 0:
            erreur = resultat.stderr.decode("utf-8", errors="replace")[-500:]
            raise RuntimeError(f"ffmpeg a échoué : {erreur}")
        with open(chemin_out, "rb") as f_out:
            return f_out.read()
    finally:
        for chemin in (chemin_in, chemin_out):
            try:
                os.remove(chemin)
            except OSError:
                pass


def _duree_fichier_secondes(chemin: str) -> float:
    """
    Renvoie la durée (en secondes) d'un fichier audio/vidéo via ffprobe.
    
    ffprobe est dans le paquet ffmpeg
    """
    resultat = subprocess.run(
        [
            "ffprobe", "-v", "error",
            "-show_entries", "format=duration",
            "-of", "default=noprint_wrappers=1:nokey=1",
            chemin,
        ],
        capture_output=True, timeout=60,
    )
    return float(resultat.stdout.decode("utf-8", errors="replace").strip() or 0.0)


# Transcrit un fichier audio/vidéo en texte via un serveur whisper.cpp local.
def outil_transcrire_audio(donnees: bytes, nom_fichier: str, conf: configparser.ConfigParser,
                            ignorer_debut_secondes: float = 0, ignorer_fin_secondes: float = 0) -> str:
    """
    Envoie un fichier audio (ou la piste audio extraite d'une vidéo) à un serveur
    whisper.cpp local, en utilisant son API native (/inference).
    
    Si le fichier est une vidéo (mp4, mkv, mov, avi, webm), sa piste audio est d'abord
    extraite et convertie en WAV mono 16 kHz via ffmpeg (whisper.cpp n'acceptant pas
    de façon fiable les conteneurs vidéo bruts, d'où l'erreur 400 côté serveur).

    Paramètres :
      donnees                  : contenu binaire brut du fichier (bytes)
      nom_fichier              : nom original du fichier (utilisé pour le type MIME et le nom envoyé)
      conf                     : configuration (section [whisper] de hermes.conf)
      ignorer_debut_secondes   : si > 0, coupe ce nombre de secondes au début avant transcription
      ignorer_fin_secondes     : si > 0, coupe ce nombre de secondes à la fin avant transcription

    Configuration attendue dans hermes.conf :
        [whisper]
        base_url          = http://localhost:8081   # racine du serveur whisper.cpp
        endpoint          = /inference              # route native whisper.cpp
        response_format   = json                    # json | text | srt | vtt
        language          = fr                      # optionnel, vide = auto-détection
    """
    try:
        # Extraction piste audio en amont (via ffmpeg système) avec découpe optionnelle début/fin
        if (Path(nom_fichier).suffix.lower() in _EXTENSIONS_VIDEO
                or ignorer_debut_secondes > 0 or ignorer_fin_secondes > 0):
            try:
                donnees = _extraire_audio_ffmpeg(
                    donnees, nom_fichier,
                    ignorer_debut_secondes=ignorer_debut_secondes,
                    ignorer_fin_secondes=ignorer_fin_secondes,
                )
                nom_fichier = Path(nom_fichier).stem + ".wav"
            except FileNotFoundError:
                return (
                    "⚠️ ffmpeg n'est pas installé (ou introuvable dans le PATH). "
                    "Installez-le pour transcrire des fichiers vidéo."
                )
            except Exception as e:
                return f"⚠️ Erreur lors du traitement audio (ffmpeg) : {e}"

        base_url = conf.get("whisper", "base_url", fallback="http://localhost:8081").rstrip("/")
        endpoint = conf.get("whisper", "endpoint", fallback="/inference")
        fmt      = conf.get("whisper", "response_format", fallback="json").strip() or "json"
        langue   = conf.get("whisper", "language", fallback="").strip()

        url          = f"{base_url}{endpoint}"
        content_type = mimetypes.guess_type(nom_fichier)[0] or "application/octet-stream"

        fichiers = {"file": (nom_fichier, donnees, content_type)}
        champs   = {"response_format": fmt}
        if langue:
            champs["language"] = langue

        reponse = requests.post(url, files=fichiers, data=champs, timeout=1200)
        reponse.raise_for_status()

        if fmt == "json":
            try:
                payload = reponse.json()
                texte = payload.get("text") or payload.get("transcription") or ""
            except ValueError:
                texte = reponse.text
        else:
            texte = reponse.text

        texte = texte.strip()
        return texte if texte else "⚠️ Transcription vide (aucun texte détecté dans l'audio)."

    except requests.exceptions.ConnectionError:
        return "⚠️ Impossible de contacter le serveur whisper.cpp "
    except requests.exceptions.Timeout:
        return "⚠️ Délai d'attente dépassé lors de la transcription audio."
    except requests.exceptions.HTTPError as e:
        return f"⚠️ Erreur HTTP du serveur whisper.cpp : {e}"
    except Exception as e:
        return f"⚠️ Erreur lors de la transcription audio : {e}"


def _extraire_id_youtube(url_ou_id: str) -> Optional[str]:
    """
    Extrait l'identifiant vidéo YouTube d'une URL
    """
    video_id = url_ou_id.strip()
    if re.match(r'^[A-Za-z0-9_-]{11}$', video_id):
        return video_id
    match = re.search(r'(?:v=|youtu\.be/|/embed/|/shorts/|/live/)([A-Za-z0-9_-]{11})', video_id)
    return match.group(1) if match else None

def _transcript_youtube_rapide(video_id: str, max_chars: int = 1024000) -> Optional[str]:
    """
    Tente de récupérer directement la transcription texte d'une vidéo YouTube via
    youtube-transcript.ai (rapide, pas de téléchargement ni de whisper.cpp nécessaire).
    Retourne None si indisponible (vidéo sans sous-titres, erreur réseau, etc.),
    auquel cas outil_transcrire_video_url se rabat sur yt-dlp + whisper.cpp.
    """
    try:
        transcript_url = f"https://youtube-transcript.ai/transcript/{video_id}.txt"
        headers = {
            "User-Agent": (
                "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36"
                "(KHTML, like Gecko) Chrome/151.0.0.0 Safari/537.36"
            )
        }
        resp = requests.get(transcript_url, headers=headers, timeout=15)
        resp.raise_for_status()
        texte = resp.text.strip()
        if not texte:
            return None
        max_chars = min(int(max_chars), 1024000)
        if len(texte) > max_chars:
            texte = texte[:max_chars] + f"\n\n[… transcription tronquée à {max_chars} caractères]"
        return texte
    except Exception:
        return None

def _url_video_autorisee(url: str, conf: configparser.ConfigParser) -> bool:
    """
    Vérifie l'URL contre la liste blanche [video_url] domaines_autorises de hermes.conf.
    Vérification faite ici (côté serveur), pas seulement dans la description du tool, pour empêcher
    le LLM (ou un prompt injecté) de faire télécharger une URL arbitraire via yt-dlp.
    "all" ou "*" dans domaines_autorises désactive la liste blanche (tous les sites autorisés).
    """
    domaines_brut = conf.get("video_url", "domaines_autorises", fallback="")
    domaines = [d.strip().lower() for d in domaines_brut.split(",") if d.strip()]
    if "all" in domaines or "*" in domaines:
        return True
    if not domaines:
        return False
    try:
        hote = urllib.parse.urlparse(url).netloc.lower()
    except ValueError:
        return False
    hote = hote.split("@")[-1].split(":")[0]  # retire d'éventuels user@ ou :port
    return any(hote == d or hote.endswith("." + d) for d in domaines)


def outil_transcrire_video_url(url: str, conf: configparser.ConfigParser,
                                ignorer_debut_secondes: float = 0, ignorer_fin_secondes: float = 0) -> str:
    """
    Transcrit en texte une vidéo hébergée en ligne, restreinte aux domaines listés dans
    hermes.conf ([video_url] domaines_autorises).

    Deux chemins possibles :
      1. Rapide : si l'URL est une vidéo YouTube et que des sous-titres sont disponibles,
         la transcription est récupérée directement via youtube-transcript.ai — pas de
         téléchargement ni de whisper.cpp nécessaire.
      2. Générique : sinon (ou si le chemin rapide échoue), la piste audio est téléchargée
         via yt-dlp, normalisée via _extraire_audio_ffmpeg (mêmes optimisations ffmpeg que
         pour un fichier audio uploadé : silences retirés, découpe début/fin, etc.), puis
         transcrite via whisper.cpp (outil_transcrire_audio).

    Paramètres :
      url                      : URL de la vidéo (doit appartenir à un domaine autorisé)
      conf                     : configuration ([whisper] et [video_url] de hermes.conf)
      ignorer_debut_secondes   : si > 0, coupe ce nombre de secondes au début avant transcription
                                 (ignoré pour le chemin rapide YouTube, qui renvoie la transcription complète)
      ignorer_fin_secondes     : si > 0, coupe ce nombre de secondes à la fin avant transcription

    Nécessite que le binaire `yt-dlp` soit installé et accessible dans le PATH (ou module python installé dans).

    Configuration attendue dans hermes.conf :
        [video_url]
        domaines_autorises = youtube.com, twitch.tv, arte.tv
    """
    if not url or not url.strip():
        return "⚠️ Aucune URL fournie."
    url = url.strip()

    if not _url_video_autorisee(url, conf):
        domaines_brut = conf.get("video_url", "domaines_autorises", fallback="")
        return (
            "⚠️ Ce site n'est pas autorisé pour la transcription vidéo. "
            f"Domaines autorisés : {domaines_brut or '(aucun configuré)'}."
        )

    # Chemin rapide : YouTube avec sous-titres disponibles
    video_id = _extraire_id_youtube(url)
    if video_id and ignorer_debut_secondes == 0 and ignorer_fin_secondes == 0:
        rapide = _transcript_youtube_rapide(video_id)
        if rapide:
            return f"📺 Transcription YouTube (ID : {video_id}) :\n\n{rapide}"
        # Pas de sous-titres ou erreur : on continue avec le chemin générique ci-dessous

    # Chemin générique : téléchargement audio (yt-dlp) puis transcription (ffmpeg + whisper.cpp)
    with tempfile.TemporaryDirectory() as tmp_dir:
        motif_sortie = os.path.join(tmp_dir, "audio.%(ext)s")
        try:
            resultat = subprocess.run(
                [
                    "yt-dlp",
                    "-x", "--audio-format", "wav", "--audio-quality", "0",
                    "-o", motif_sortie,
                    url,
                ],
                capture_output=True, timeout=1800,
            )
        except FileNotFoundError:
            return "⚠️ yt-dlp n'est pas installé (ou introuvable dans le PATH)."
        except subprocess.TimeoutExpired:
            return "⚠️ Délai d'attente dépassé lors du téléchargement de la vidéo (yt-dlp)."

        if resultat.returncode != 0:
            erreur = resultat.stderr.decode("utf-8", errors="replace")[-500:]
            return f"⚠️ Échec du téléchargement de la vidéo (yt-dlp) : {erreur}"

        fichiers = [f for f in os.listdir(tmp_dir) if f.startswith("audio.")]
        if not fichiers:
            return "⚠️ yt-dlp n'a produit aucun fichier audio exploitable."
        nom_fichier = fichiers[0]
        chemin_audio = os.path.join(tmp_dir, nom_fichier)

        try:
            with open(chemin_audio, "rb") as f:
                donnees = f.read()
        except OSError as e:
            return f"⚠️ Erreur lors de la lecture du fichier audio téléchargé : {e}"

        # Normalisation via _extraire_audio_ffmpeg (silences retirés, découpe début/fin, mono 16 kHz)
        try:
            donnees = _extraire_audio_ffmpeg(
                donnees, nom_fichier,
                ignorer_debut_secondes=ignorer_debut_secondes,
                ignorer_fin_secondes=ignorer_fin_secondes,
            )
            nom_fichier = Path(nom_fichier).stem + ".wav"
        except FileNotFoundError:
            return "⚠️ ffmpeg n'est pas installé (ou introuvable dans le PATH)."
        except Exception as e:
            return f"⚠️ Erreur lors du traitement audio (ffmpeg) : {e}"

        # Audio déjà normalisé/découpé : on passe ignorer_debut/fin=0 pour éviter un second passage ffmpeg
        return outil_transcrire_audio(donnees, nom_fichier, conf)


# Génère une image à partir d'un prompt texte via un serveur stablediffusion.cpp local.
def outil_generation_image(prompt: str, conf: configparser.ConfigParser,
                            negative_prompt: str = "", taille: str = "") -> str:
    """
    Envoie un prompt texte à un serveur stablediffusion.cpp local (API compatible
    OpenAI /v1/images/generations) et récupère l'image générée encodée en base64.

    Paramètres :
      prompt          : description de l'image à générer
      conf            : configuration (section [image] de hermes.conf)
      negative_prompt : éléments à éviter dans l'image (optionnel)
      taille          : dimensions "LARGEURxHAUTEUR" (optionnel, sinon valeur de la config)

    Configuration attendue dans hermes.conf :
        [image]
        base_url = http://localhost:8082           # racine du serveur stablediffusion.cpp
        endpoint = /v1/images/generations           # route compatible OpenAI
        size     = 512x512                          # taille par défaut
        steps    =                                  # nombre d'étapes (vide = défaut serveur)

    Retourne un JSON avec {b64, mime, nom, format, __image_generee__: true}.
    """
    import json as _json

    if conf is None:
        conf = configparser.ConfigParser()

    try:
        base_url = conf.get("image", "base_url", fallback="http://localhost:8082").rstrip("/")
        endpoint = conf.get("image", "endpoint", fallback="/v1/images/generations")
        taille   = taille or conf.get("image", "size", fallback="512x512")
        steps    = conf.get("image", "steps", fallback="").strip()

        url     = f"{base_url}{endpoint}"
        payload = {
            "prompt":          prompt,
            "size":            taille,
            "n":               1,
            "response_format": "b64_json",
        }
        if negative_prompt:
            payload["negative_prompt"] = negative_prompt
        if steps:
            payload["steps"] = int(steps)

        reponse = requests.post(url, json=payload, timeout=600)
        reponse.raise_for_status()
        donnees = reponse.json()

        images = donnees.get("data") or []
        if not images or "b64_json" not in images[0]:
            return "⚠️ Le serveur de génération d'image n'a retourné aucune image."

        b64 = images[0]["b64_json"]
        nom = f"image_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"

        return _json.dumps({
            "__image_generee__": True,
            "b64":    b64,
            "mime":   "image/png",
            "nom":    nom,
            "format": "png",
        })

    except requests.exceptions.ConnectionError:
        return "⚠️ Impossible de contacter le serveur stablediffusion.cpp."
    except requests.exceptions.Timeout:
        return "⚠️ Délai d'attente dépassé lors de la génération d'image."
    except requests.exceptions.HTTPError as e:
        return f"⚠️ Erreur HTTP du serveur stablediffusion.cpp : {e}"
    except Exception as e:
        return f"⚠️ Erreur lors de la génération d'image : {e}"


# Retourne la date et l'heure actuelles.
def outil_datetime() -> str:
    """
    Retourne la date et l'heure locales formatées en français.
    Aucun paramètre requis.
    """
    return datetime.now().strftime("Date : %A %d %B %Y — Heure : %H:%M:%S")


# Génère un fichier (txt, md, pdf, docx, xlsx) et retourne son contenu en base64.
def outil_generer_fichier(contenu: str, format: str, nom_fichier: str) -> str:
    """
    Génère un fichier téléchargeable dans le format demandé.

    Paramètres :
      contenu      : texte à mettre dans le fichier (markdown accepté)
      format       : extension cible — txt | md | pdf | docx | xlsx
      nom_fichier  : nom du fichier sans extension (ex : rapport_meteo)
    Retourne un JSON avec {b64, mime, nom, format, __fichier_genere__: true}.
    """
    import io, base64 as _b64, json as _json

    fmt = format.lower().lstrip(".")

    # Nom propre
    if "." not in nom_fichier:
        nom_fichier = f"{nom_fichier}.{fmt}"

    try:
        # TXT / MD
        if fmt in ("txt", "md"):
            data = contenu.encode("utf-8")
            mime = "text/plain" if fmt == "txt" else "text/markdown"

        # PDF (fpdf2)
        elif fmt == "pdf":
            from fpdf import FPDF
            pdf = FPDF()
            pdf.set_margins(15, 15, 15)
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()

            for line in contenu.split("\n"):
                stripped = line.rstrip()
                if stripped == "":
                    pdf.ln(4)
                elif stripped.startswith("### "):
                    pdf.set_font("Helvetica", "B", 12)
                    pdf.multi_cell(0, 7, stripped[4:])
                    pdf.set_font("Helvetica", size=11)
                elif stripped.startswith("## "):
                    pdf.set_font("Helvetica", "B", 14)
                    pdf.multi_cell(0, 8, stripped[3:])
                    pdf.set_font("Helvetica", size=11)
                elif stripped.startswith("# "):
                    pdf.set_font("Helvetica", "B", 16)
                    pdf.multi_cell(0, 10, stripped[2:])
                    pdf.set_font("Helvetica", size=11)
                elif stripped.startswith(("- ", "* ", "• ")):
                    pdf.set_font("Helvetica", size=11)
                    pdf.multi_cell(0, 6, "  • " + stripped[2:])
                else:
                    pdf.set_font("Helvetica", size=11)
                    pdf.multi_cell(0, 6, stripped)

            data = bytes(pdf.output())
            mime = "application/pdf"

        # DOCX (python-docx)
        elif fmt == "docx":
            from docx import Document
            from docx.shared import Pt

            doc = Document()
            for line in contenu.split("\n"):
                stripped = line.rstrip()
                if stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith(("- ", "* ", "• ")):
                    doc.add_paragraph(stripped[2:], style="List Bullet")
                elif stripped == "":
                    doc.add_paragraph("")
                else:
                    doc.add_paragraph(stripped)

            buf = io.BytesIO()
            doc.save(buf)
            data = buf.getvalue()
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        # XLSX (openpyxl)
        elif fmt == "xlsx":
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment

            wb = openpyxl.Workbook()
            ws = wb.active
            lines = [l for l in contenu.split("\n")]

            # Détection tableau markdown (lignes avec |)
            table_lines = [l for l in lines if "|" in l and l.strip().startswith("|")]

            if table_lines:
                row_idx = 1
                header_done = False
                for line in lines:
                    if not line.strip():
                        continue
                    if "|" not in line:
                        continue
                    # Ligne séparatrice |---|---|
                    if all(c in "-|: " for c in line):
                        header_done = True
                        continue
                    cells = [c.strip() for c in line.strip().strip("|").split("|")]
                    for col_idx, cell in enumerate(cells, 1):
                        c = ws.cell(row=row_idx, column=col_idx, value=cell)
                        if not header_done:  # ligne d'en-tête
                            c.font = Font(bold=True)
                            c.fill = PatternFill("solid", fgColor="4472C4")
                            c.font = Font(bold=True, color="FFFFFF")
                        c.alignment = Alignment(wrap_text=True)
                    row_idx += 1
                # Auto-largeur approximative
                for col in ws.columns:
                    max_len = max((len(str(c.value or "")) for c in col), default=10)
                    ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)
            else:
                # Texte brut : une ligne = une cellule A?
                for row_idx, line in enumerate(lines, 1):
                    ws.cell(row=row_idx, column=1, value=line)

            buf = io.BytesIO()
            wb.save(buf)
            data = buf.getvalue()
            mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

        else:
            return f"⚠️ Format « {format} » non supporté. Formats disponibles : txt, md, pdf, docx, xlsx"

        b64 = _b64.b64encode(data).decode("utf-8")
        return _json.dumps({
            "__fichier_genere__": True,
            "b64":    b64,
            "mime":   mime,
            "nom":    nom_fichier,
            "format": fmt,
        })

    except ImportError as e:
        return f"⚠️ Bibliothèque manquante pour générer le fichier .{fmt} : {e}"
    except Exception as e:
        return f"⚠️ Erreur lors de la génération du fichier .{fmt} : {e}"


# Catalogue JSON (schéma pour le LLM
CATALOGUE_OUTILS = [
    {
        "type": "function",
        "function": {
            "name": "outil_meteo",
            "description": "Donne la météo complète et en temps réel d'une ville (température, humidité, vent, état du ciel).",
            "parameters": {
                "type": "object",
                "properties": {
                    "ville": {
                        "type": "string",
                        "description": "Nom de la ville (ex : Paris, London, Tokyo).",
                    }
                },
                "required": ["ville"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "outil_wiki",
            "description": "Recherche et retourne un résumé Wikipédia sur n'importe quel sujet.",
            "parameters": {
                "type": "object",
                "properties": {
                    "sujet": {
                        "type": "string",
                        "description": "Sujet à rechercher sur Wikipédia (ex : Tour Eiffel, Albert Einstein).",
                    }
                },
                "required": ["sujet"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "outil_argent",
            "description": "Convertit un montant d'une devise vers une autre avec les taux en temps réel.",
            "parameters": {
                "type": "object",
                "properties": {
                    "montant":      {"type": "number", "description": "Valeur à convertir (ex : 100)."},
                    "de_monnaie":   {"type": "string", "description": "Code ISO 4217 source (ex : EUR)."},
                    "vers_monnaie": {"type": "string", "description": "Code ISO 4217 cible  (ex : USD)."},
                },
                "required": ["montant", "de_monnaie", "vers_monnaie"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "outil_duckduckgo",
            "description": (
                "Effectue une recherche web via DuckDuckGo et retourne les N premiers "
                "résultats (titre, URL, extrait). Utilise cet outil pour répondre à des "
                "questions nécessitant des informations récentes ou factuelles."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "La requête de recherche en langage naturel.",
                    },
                    "num_results": {
                        "type": "integer",
                        "description": "Nombre de résultats à retourner (défaut 5, max 10).",
                        "default": 5,
                    },
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "outil_recup_page",
            "description": (
                "Télécharge et extrait le texte principal d'une page web à partir "
                "de son URL. Utile pour lire le contenu complet d'un article ou d'une page."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "L'URL complète de la page à récupérer.",
                    },
                    "max_chars": {
                        "type": "integer",
                        "description": "Nombre maximum de caractères à retourner (défaut 1024000).",
                        "default": 1024000,
                    },
                },
                "required": ["url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "outil_datetime",
            "description": "Retourne la date et l'heure actuelles (horloge locale du serveur).",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "outil_generer_fichier",
            "description": (
                "Génère un fichier téléchargeable (txt, md, pdf, docx, xlsx) à partir d'un contenu textuel. "
                "Utilise cet outil quand l'utilisateur demande explicitement à exporter, sauvegarder, "
                "télécharger ou créer un fichier avec le résultat d'une réponse."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "contenu": {
                        "type": "string",
                        "description": "Contenu complet à mettre dans le fichier (texte brut ou markdown).",
                    },
                    "format": {
                        "type": "string",
                        "enum": ["txt", "md", "pdf", "docx", "xlsx"],
                        "description": "Format du fichier à générer.",
                    },
                    "nom_fichier": {
                        "type": "string",
                        "description": "Nom du fichier sans extension (ex : rapport_meteo, synthese_2024).",
                    },
                },
                "required": ["contenu", "format", "nom_fichier"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "outil_generation_image",
            "description": (
                "Génère une image à partir d'une description textuelle (prompt) via un serveur "
                "stablediffusion.cpp local. Utilise cet outil quand l'utilisateur demande de "
                "créer, dessiner, générer ou illustrer une image."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {
                        "type": "string",
                        "description": (
                            "Description détaillée de l'image à générer "
                            "(en anglais de préférence pour de meilleurs résultats)."
                        ),
                    },
                    "negative_prompt": {
                        "type": "string",
                        "description": "Éléments à éviter dans l'image générée (optionnel).",
                    },
                    "taille": {
                        "type": "string",
                        "description": (
                            "Dimensions de l'image au format LARGEURxHAUTEUR (ex : 512x512). "
                            "Optionnel, valeur par défaut définie dans la config."
                        ),
                    },
                },
                "required": ["prompt"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "outil_transcrire_audio",
            "description": (
                "Transcrit en texte le fichier audio ou vidéo joint au message courant "
                "(via un serveur whisper.cpp local). À utiliser uniquement lorsque l'utilisateur "
                "a joint un fichier audio/vidéo et souhaite en connaître le contenu dicté "
                "(résumé, réponse à une question sur l'enregistrement, etc.). "
                "Ne fonctionne que s'il y a effectivement un fichier audio/vidéo joint. "
                "Si l'utilisateur demande de retirer une partie du début ou de la fin avec par exemple "
                "« enlève les 5 premières minutes », « saute les 2 premières minutes », "
                "« coupe la dernière minute »), convertis sa demande en secondes et renseigne "
                "ignorer_debut_secondes et/ou ignorer_fin_secondes en conséquence."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "ignorer_debut_secondes": {
                        "type": "number",
                        "description": (
                            "Nombre de secondes à couper au début du fichier avant transcription "
                            "(0 par défaut). Ex : l'utilisateur dit « enlève les 5 premières minutes » -> 300."
                        ),
                    },
                    "ignorer_fin_secondes": {
                        "type": "number",
                        "description": (
                            "Nombre de secondes à couper à la fin du fichier avant transcription "
                            "(0 par défaut). Ex : l'utilisateur dit « enlève la dernière minute » -> 60."
                        ),
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "outil_transcrire_video_url",
            "description": (
                "Récupère le contenu textuel (transcription) d'une vidéo en ligne à partir de son URL. "
                "IMPORTANT : le modèle n'a AUCUN moyen de connaître le contenu d'une vidéo à partir de "
                "sa seule URL, dès qu'un message contient un lien vers une vidéo (YouTube, Twitch, "
                "Dailymotion, Arte, etc., peu importe qu'il soit collé seul ou dans une phrase), cet outil "
                "DOIT être appelé en premier pour en récupérer le contenu, AVANT de répondre à la demande "
                "de l'utilisateur (résumer, analyser, traduire, répondre à une question, extraire des "
                "infos...), même si l'utilisateur ne mentionne pas explicitement les mots « transcrire » "
                "ou « extraire ». Ne jamais deviner ou halluciner le contenu d'une vidéo à partir de son URL. "
                "Pour YouTube, l'outil utilise en priorité la transcription/les sous-titres officiels quand "
                "disponibles (rapide) ; sinon, il télécharge la piste audio et la transcrit via whisper.cpp. "
                "Si l'utilisateur demande de retirer une partie du début ou de la fin (par ex. "
                "« enlève les 5 premières minutes »), convertis sa demande en secondes et renseigne "
                "ignorer_debut_secondes et/ou ignorer_fin_secondes en conséquence."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "URL de la vidéo (doit appartenir à un domaine autorisé par le serveur, ex : YouTube, Twitch, Arte).",
                    },
                    "ignorer_debut_secondes": {
                        "type": "number",
                        "description": "Secondes à couper au début avant transcription (0 par défaut).",
                    },
                    "ignorer_fin_secondes": {
                        "type": "number",
                        "description": "Secondes à couper à la fin avant transcription (0 par défaut).",
                    },
                },
                "required": ["url"],
            },
        },
    },
]

# Emojis d'affichage (utilisés dans l'interface web)
ICONES_OUTILS = {
    "outil_meteo":              "🌤️ Météo",
    "outil_wiki":               "📖 Wikipédia",
    "outil_argent":             "💱 Change",
    "outil_duckduckgo":         "🔍 Recherche web",
    "outil_recup_page":         "🌐 Lecture page",
    "outil_datetime":           "🕐 Date & Heure",
    "outil_generer_fichier":    "💾 Génération fichier",
    "outil_transcrire_audio":   "🎙️ Transcription audio/vidéo",
    "outil_transcrire_video_url": "📡 Transcription vidéo (URL)",
    "outil_generation_image":   "🎨 Génération d'image",
}



# Fonctions utilitaires
# Retourne la liste des outils activés selon la section [tools] du .conf. Permet d'activer/désactiver chaque outil sans toucher au code.
def outils_actifs(conf: configparser.ConfigParser) -> list:
    mapping = {
        "enable_meteo":             "outil_meteo",
        "enable_wiki":              "outil_wiki",
        "enable_argent":            "outil_argent",
        "enable_duckduckgo":        "outil_duckduckgo",
        "enable_recup_page":        "outil_recup_page",
        "enable_datetime":          "outil_datetime",
        "enable_generer_fichier":   "outil_generer_fichier",
        "enable_transcrire_audio":  "outil_transcrire_audio",
        "enable_transcrire_video_url": "outil_transcrire_video_url",
        "enable_generation_image": "outil_generation_image",
    }
    actifs = []
    for cle, nom in mapping.items():
        if conf.getboolean("tools", cle, fallback=True):
            actifs.extend([o for o in CATALOGUE_OUTILS if o["function"]["name"] == nom])
    return actifs

# Dispatcher central : appelle la bonne fonction selon le nom renvoyé par le LLM.
# Toujours utiliser cette fonction plutôt qu'appeler les outils directement.
def executer_outil(nom: str, args: dict, conf: configparser.ConfigParser = None) -> str:
    if nom == "outil_meteo":
        return outil_meteo(args["ville"])
    elif nom == "outil_wiki":
        return outil_wiki(args["sujet"])
    elif nom == "outil_argent":
        return outil_argent(args["montant"], args["de_monnaie"], args["vers_monnaie"])
    elif nom == "outil_duckduckgo":
        return outil_duckduckgo(args["query"], args.get("num_results", 10))
    elif nom == "outil_recup_page":
        return outil_recup_page(args["url"], args.get("max_chars", 1024000))
    elif nom == "outil_datetime":
        return outil_datetime()
    elif nom == "outil_generer_fichier":
        return outil_generer_fichier(args["contenu"], args["format"], args["nom_fichier"])
    elif nom == "outil_generation_image":
        return outil_generation_image(args["prompt"], conf, args.get("negative_prompt", ""), args.get("taille", ""))
    elif nom == "outil_transcrire_audio":
        # Intercepté en amiont par hermes-web, ce cas est déclencé si jamais on utilise l'outil depuis la CLI
        return "⚠️ Aucun fichier audio/vidéo joint : cet outil nécessite un fichier attaché depuis l'interface web."
    elif nom == "outil_transcrire_video_url":
        return outil_transcrire_video_url(
            args["url"], conf,
            ignorer_debut_secondes=args.get("ignorer_debut_secondes", 0) or 0,
            ignorer_fin_secondes=args.get("ignorer_fin_secondes", 0) or 0,
        )
    return f"⚠️ Outil inconnu : « {nom} »."
